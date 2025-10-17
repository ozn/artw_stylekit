"""DOCX document builder."""
from pathlib import Path
from typing import Dict, List
import json
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from ..logger import logger

class DocxBuilder:
    """Build DOCX documents from outline."""
    
    def __init__(self):
        self.doc = Document()
        self._setup_styles()
    
    def _setup_styles(self):
        """Setup document styles."""
        # Normal style
        style = self.doc.styles['Normal']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
        # Heading styles
        for i in range(1, 4):
            style = self.doc.styles[f'Heading {i}']
            font = style.font
            font.name = 'Times New Roman'
            font.size = Pt(14 - i)
            font.bold = True
    
    def build_from_outline(self, outline: Dict) -> Document:
        """
        Build document from outline JSON.
        
        Args:
            outline: Article outline dictionary
            
        Returns:
            Document object
        """
        # Title
        title = self.doc.add_heading(outline.get('title', 'Başlık'), 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Abstract (TR)
        self.doc.add_heading('Özet', 1)
        self.doc.add_paragraph(outline.get('abstract_tr', ''))
        
        # Keywords (TR)
        keywords_tr = ', '.join(outline.get('keywords_tr', []))
        p = self.doc.add_paragraph()
        p.add_run('Anahtar Kelimeler: ').bold = True
        p.add_run(keywords_tr).italic = True
        
        self.doc.add_paragraph()  # Spacing
        
        # Abstract (EN)
        self.doc.add_heading('Abstract', 1)
        self.doc.add_paragraph(outline.get('abstract_en', ''))
        
        # Keywords (EN)
        keywords_en = ', '.join(outline.get('keywords_en', []))
        p = self.doc.add_paragraph()
        p.add_run('Keywords: ').bold = True
        p.add_run(keywords_en).italic = True
        
        self.doc.add_page_break()
        
        # Sections
        for section in outline.get('sections', []):
            self._add_section(section)
        
        # References
        self._add_references(outline.get('references', []))
        
        return self.doc
    
    def _add_section(self, section: Dict):
        """Add a section to document."""
        # Section title
        self.doc.add_heading(section.get('title', ''), 1)
        
        # Placeholder content
        content = section.get('content', '')
        if not content:
            # Generate placeholder
            key_points = section.get('key_points', [])
            content = f"[Bu bölüm yazılacak. Tahmini: {section.get('estimated_words', 500)} kelime]\n\n"
            content += "Ana noktalar:\n"
            for point in key_points:
                content += f"- {point}\n"
            content += f"\nMinimum atıf sayısı: {section.get('required_citations', 2)}"
        
        self.doc.add_paragraph(content)
        
        # Subsections
        for subsection in section.get('subsections', []):
            if isinstance(subsection, str):
                self.doc.add_heading(subsection, 2)
                self.doc.add_paragraph('[Alt bölüm içeriği buraya gelecek]')
            elif isinstance(subsection, dict):
                self.doc.add_heading(subsection.get('title', ''), 2)
                self.doc.add_paragraph(subsection.get('content', '[İçerik buraya gelecek]'))
    
    def _add_references(self, references: List):
        """Add references section."""
        self.doc.add_page_break()
        self.doc.add_heading('Kaynakça', 1)
        
        if not references:
            self.doc.add_paragraph('[Kaynaklar buraya eklenecek]')
            return
        
        for ref in references:
            if isinstance(ref, str):
                self.doc.add_paragraph(ref, style='Normal')
            elif isinstance(ref, dict):
                citation = ref.get('apa_citation', ref.get('citation', ''))
                self.doc.add_paragraph(citation, style='Normal')
    
    def save(self, path: Path):
        """Save document to file."""
        path.parent.mkdir(parents=True, exist_ok=True)
        self.doc.save(str(path))
        logger.info(f"Document saved to {path}")
